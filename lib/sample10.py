import requests
from docx import Document

def get_stock_price(symbol):
    url = f"https://api.example.com/stocks/{symbol}"
    response = requests.get(url)
    data = response.json()

    if response.status_code == 404:
        return None

    price = data["price"]
    return price

def save_prices_to_docx(prices, filename):
    doc = Document()

    doc.add_heading("Stock Prices", level=1)

    for symbol, price in prices.items():
        doc.add_paragraph(f"{symbol}: {price}")

    doc.save(filename)
    print(f"Сохранено в {filename}")

def main():
    print("Виджет биржевых цен")

    symbols = ["AAPL", "GOOGL", "MSFT"]  # Пример символов акций

    prices = {}

    for symbol in symbols:
        price = get_stock_price(symbol)
        if price:
            prices[symbol] = price

    if prices:
        filename = "stock_prices.docx"
        save_prices_to_docx(prices, filename)
    else:
        print("Не удалось получить данные о ценах акций.")

    print("Спасибо за использование виджета биржевых цен. До свидания!")

if __name__ == "__main__":
    main()
